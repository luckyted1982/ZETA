"""
文档解析模块
支持 PDF、DOCX、XLSX、TXT、Markdown 等格式
"""

import os
from typing import Optional, Dict, Any

# 第三方文档解析库
from PyPDF2 import PdfReader  # PDF 文本提取
from docx import Document  # Word 文档解析
from openpyxl import load_workbook  # Excel 表格解析


class DocumentParser:
    """文档解析器

    提供统一的文档解析入口，支持多种常见格式的文本提取和元信息提取。
    """

    MAX_TEXT_LENGTH = 50000  # 最大文本长度限制，防止超大文档导致内存溢出

    @classmethod
    def parse(
        cls, file_path: str, file_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """解析文档，返回文本内容和元信息

        Args:
            file_path: 文件在本地存储的绝对路径
            file_type: 文件类型（如 pdf/docx/xlsx），若为空则自动检测扩展名

        Returns:
            字典结构：
            {
                "text": str,        # 解析后的纯文本（已截断至 MAX_TEXT_LENGTH）
                "summary": str,     # 前 500 字的摘要
                "file_type": str,   # 文件类型
                "pages": int,       # 页数/工作表数/段落估算数
                "status": "success" | "error",
                "error": str | None,
            }
        """
        if not file_type:
            file_type = cls._detect_type(file_path)

        try:
            if file_type == "pdf":
                result = cls._parse_pdf(file_path)
            elif file_type == "docx":
                result = cls._parse_docx(file_path)
            elif file_type == "xlsx":
                result = cls._parse_xlsx(file_path)
            elif file_type in ("txt", "md", "markdown"):
                result = cls._parse_text(file_path)
            else:
                return {
                    "status": "error",
                    "error": f"不支持的文件类型: {file_type}",
                    "text": "",
                    "summary": "",
                    "file_type": file_type,
                    "pages": 0,
                }

            # 截取摘要：取前 500 字，若超出则追加省略号
            text = result.get("text", "")
            result["summary"] = text[:500] + ("..." if len(text) > 500 else "")
            result["file_type"] = file_type
            result["status"] = "success"
            result["error"] = None

            return result

        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "text": "",
                "summary": "",
                "file_type": file_type,
                "pages": 0,
            }

    @classmethod
    def _detect_type(cls, file_path: str) -> str:
        """根据文件扩展名检测文件类型

        Args:
            file_path: 文件路径

        Returns:
            文件类型标识字符串（如 pdf、docx、xlsx、txt 等）
        """
        ext = os.path.splitext(file_path)[1].lower().lstrip(".")
        type_map = {
            "pdf": "pdf",
            "docx": "docx",
            "xlsx": "xlsx",
            "xls": "xlsx",  # 旧版 Excel 统一按 xlsx 处理
            "txt": "txt",
            "md": "md",
            "markdown": "md",
        }
        return type_map.get(ext, "unknown")

    @classmethod
    def _parse_pdf(cls, file_path: str) -> Dict[str, Any]:
        """解析 PDF 文件

        Args:
            file_path: PDF 文件路径

        Returns:
            包含 text（文本）和 pages（页数）的字典
        """
        reader = PdfReader(file_path)
        pages_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
        full_text = "\n\n".join(pages_text)
        return {
            "text": full_text[: cls.MAX_TEXT_LENGTH],
            "pages": len(reader.pages),
        }

    @classmethod
    def _parse_docx(cls, file_path: str) -> Dict[str, Any]:
        """解析 Word 文档（DOCX）

        Args:
            file_path: DOCX 文件路径

        Returns:
            包含 text（文本）和 pages（估算页数）的字典
        """
        doc = Document(file_path)
        # 提取段落文本
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        # 额外提取表格内容，按 " | " 拼接单元格
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells if cell.text.strip()]
                if cells:
                    full_text += "\n" + " | ".join(cells)
        # 估算页数：按平均每页 30 个段落粗略估算
        return {
            "text": full_text[: cls.MAX_TEXT_LENGTH],
            "pages": len(doc.paragraphs) // 30 + 1,
        }

    @classmethod
    def _parse_xlsx(cls, file_path: str) -> Dict[str, Any]:
        """解析 Excel 表格（XLSX/XLS）

        Args:
            file_path: Excel 文件路径

        Returns:
            包含 text（文本）和 pages（工作表数）的字典
        """
        wb = load_workbook(file_path, data_only=True)
        all_text = []
        for sheet in wb.worksheets:
            all_text.append(f"=== 工作表: {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    all_text.append(" | ".join(cells))
        full_text = "\n".join(all_text)
        return {
            "text": full_text[: cls.MAX_TEXT_LENGTH],
            "pages": len(wb.worksheets),
        }

    @classmethod
    def _parse_text(cls, file_path: str) -> Dict[str, Any]:
        """解析纯文本文件（TXT / Markdown）

        Args:
            file_path: 文本文件路径

        Returns:
            包含 text（文本）和 pages（估算页数）的字典
        """
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        # 估算页数：按平均每页 50 行粗略估算
        return {
            "text": text[: cls.MAX_TEXT_LENGTH],
            "pages": text.count("\n") // 50 + 1,
        }


def parse_document(file_path: str, file_type: Optional[str] = None) -> Dict[str, Any]:
    """便捷函数：解析单个文档

    Args:
        file_path: 文件路径
        file_type: 可选的文件类型，为空则自动识别

    Returns:
        解析结果字典，结构与 DocumentParser.parse 一致
    """
    return DocumentParser.parse(file_path, file_type)
